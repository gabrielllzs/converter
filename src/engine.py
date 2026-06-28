import os
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx2pdf import convert as docx_to_pdf_compiler
from PIL import Image, ImageOps

class ConverterEngine:
    def convert_to_word(self, file_paths, output_folder):
        successful = []
        for path in file_paths:
            filename = os.path.basename(path)
            name, ext = os.path.splitext(filename)
            docx_path = os.path.join(output_folder, f"{name}.docx")
            doc = Document()
            
            # Setup margins
            for section in doc.sections:
                section.top_margin = Inches(0.25)
                section.bottom_margin = Inches(0.25)
                section.left_margin = Inches(0.25)
                section.right_margin = Inches(0.25)
                section.page_width = Inches(8.5)
                section.page_height = Inches(11.0)
                
            target_width = 7.8
            target_max_height = 9.8
            max_ratio = target_max_height / target_width
            
            if ext.lower() == '.pdf':
                pdf_doc = fitz.open(path)
                full_text = "".join([page.get_text() for page in pdf_doc])
                
                if full_text.strip():
                    for section in doc.sections:
                        section.top_margin = section.bottom_margin = Inches(1.0)
                        section.left_margin = section.right_margin = Inches(1.0)
                    for paragraph in full_text.split('\n'):
                        if paragraph.strip(): doc.add_paragraph(paragraph)
                else:
                    is_first = True
                    for page in pdf_doc:
                        rect = page.rect
                        if rect.height / rect.width > max_ratio:
                            slice_height = rect.width * max_ratio
                            y = 0
                            while y < rect.height:
                                if not is_first: doc.add_page_break()
                                is_first = False
                                y_next = min(y + slice_height, rect.height)
                                clip = fitz.Rect(0, y, rect.width, y_next)
                                pix = page.get_pixmap(clip=clip, dpi=200)
                                doc.add_picture(io.BytesIO(pix.tobytes("png")), width=Inches(target_width))
                                y += slice_height
                        else:
                            if not is_first: doc.add_page_break()
                            is_first = False
                            pix = page.get_pixmap(dpi=200)
                            doc.add_picture(io.BytesIO(pix.tobytes("png")), width=Inches(target_width))
                pdf_doc.close()
            
            elif ext.lower() in ['.png', '.jpg', '.jpeg']:
                img_doc = fitz.open(path)
                page = img_doc[0]
                pix = page.get_pixmap(dpi=200)
                doc.add_picture(io.BytesIO(pix.tobytes("png")), width=Inches(target_width))
                img_doc.close()
                
            doc.save(docx_path)
            successful.append(docx_path)
        return successful

    def convert_to_pdf(self, file_paths, output_folder):
        successful = []
        for path in file_paths:
            filename = os.path.basename(path)
            name, ext = os.path.splitext(filename)
            pdf_path = os.path.join(output_folder, f"{name}.pdf")
        
        if ext.lower() == '.docx':
                docx_to_pdf_compiler(path, pdf_path)
                successful.append(pdf_path)
                
        elif ext.lower() in ['.png', '.jpg', '.jpeg']:
            img = Image.open(path)
                
            # Fix for RGBA (transparency) images: convert to RGB before saving as PDF
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
                
            img.save(pdf_path, "PDF", resolution=100.0)
            successful.append(pdf_path)
            
        return successful